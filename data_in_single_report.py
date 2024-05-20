import pandas as pd
from docx import Document
from docx.enum.text import WD_ALIGN_PARAGRAPH
import os
import copy

# def excel_to_word_to_pdf(excel_path, template_path, output_dir):
#     # Read Excel file
#     try:
#         # Read Excel file, skipping the header row
#         df = pd.read_excel(excel_path, header=1)
#     except Exception as e:
#         print(f"Error reading Excel file: {e}")
#         return

#     if not os.path.exists(output_dir):
#         os.makedirs(output_dir)


#     # Define a mapping of table headers to Excel columns
#     table_mapping = {
#         "Employee Name": "Employee Name",
#         "Employee ID": "Employee Code",
#         "Emp. Job Title": "Current Position",
#     }

    
#     # Iterate through each row in the DataFrame
#     for index, row in df.iterrows():

#         try:

#             # Create a new document based on the template
#             doc = Document(template_path)

#             # Iterate over the tables in the document
#             for table in doc.tables:
#                 for row_idx, table_row in enumerate(table.rows):
#                     for cell_idx, cell in enumerate(table_row.cells):
#                         # Check if the cell text matches one of the headers
#                         cell_text = cell.text.strip()
#                         if cell_text in table_mapping:
#                             # Get the corresponding column name
#                             col_name = table_mapping[cell_text]
#                             # Get the value from the DataFrame
#                             value = str(row[col_name])
#                             # Set the value in the next cell
#                             cell_to_update = table.cell(row_idx, cell_idx + 1)
#                             cell_to_update.text = value

#                             # Center align the text
#                             for paragraph in cell_to_update.paragraphs:
#                                 paragraph.alignment = WD_ALIGN_PARAGRAPH.CENTER

#                             # If it's the Employee ID, make the text bold
#                             if cell_text == "Employee ID":
#                                 for paragraph in cell_to_update.paragraphs:
#                                     for run in paragraph.runs:
#                                         run.bold = True

#             # Save the modified Word document
#             word_filename = os.path.join(output_dir, f'report_{index + 1}.docx')
#             doc.save(word_filename)


#             print(f'Processed data for row {index + 2}')
#         except Exception as e:
#             print(f"Error processing row {index + 2}: {e}")

    

# # Define the paths
# excel_path = '/home/exotic/Downloads/PCT-3 Engineering.xlsx'
# template_path = '/home/exotic/Downloads/PCT_HR_4002_F02 Training Evaluation Form_02.docx'
# output_dir = '/home/exotic/Downloads/Reports'

# # Run the function
# excel_to_word_to_pdf(excel_path, template_path, output_dir)



# from docx import Document
# from docxcompose.composer import Composer
# from docx.enum.text import WD_BREAK
# import os

def add_page_break(combined_doc):
    # Add a paragraph and a page break to the combined document
    paragraph = combined_doc.add_paragraph()
    run = paragraph.add_run()
    run.add_break(WD_BREAK.PAGE)



# from docx import Document
# from docxcompose.composer import Composer
# from docx.enum.text import WD_BREAK
# from docx.shared import Inches
# import os

# def combine_docs(input_dir, output_file):
#     # Initialize combined_doc with content from the first document
#     first_doc_path = os.path.join(input_dir, sorted(os.listdir(input_dir))[0])
    
#     combined_doc = Document(first_doc_path)

#     # Copy the headers and footers from the initial document
#     # /home/exotic/Downloads/Reports/report_1.docx
#     composer = Composer(combined_doc)

#     # Append content from subsequent documents, retaining the initial header and footer
#     file_names = sorted(os.listdir(input_dir))
#     for index, file_name in enumerate(file_names):
#         if file_name.endswith('.docx'):
#             file_path = os.path.join(input_dir, file_name)
#             if file_path != first_doc_path:
#                 sub_doc = Document(file_path)
#                 if index > 0:
#                     add_page_break(combined_doc)
                
#                 footer_paragraphs = combined_doc.sections[0].footer.paragraphs
#                 sub_doc_footer = sub_doc.sections[0].footer
#                 for paragraph in footer_paragraphs:
#                     sub_doc_paragraph = sub_doc_footer.add_paragraph(paragraph.text)
#                     sub_doc_paragraph.style = paragraph.style
#                 composer.append(sub_doc)

#     composer.save(output_file)
#     print(f'Saved combined Word document: {output_file}')

# # Paths
# input_dir = '/home/exotic/Downloads/Reports'
# output_file = '/home/exotic/Downloads/Reports/Combined_Report.docx'

# combine_docs(input_dir, output_file)


from docx import Document
from docx.oxml.ns import qn
from docx.oxml import OxmlElement

def footer_exists(section):
    """
    Check if a footer exists in the section.
    """
    for paragraph in section.footer.paragraphs:
        if paragraph.text.strip():
            return True
    for table in section.footer.tables:
        if any(cell.text.strip() for row in table.rows for cell in row.cells):
            return True
    return False

def copy_table(source_table, target_footer):
    """
    Copy a table from the source table to the target footer.
    """
    target_table = target_footer.add_table(rows=source_table.rows, cols=source_table.columns)
    target_table.style = source_table.style

    for row_idx, row in enumerate(source_table.rows):
        for col_idx, cell in enumerate(row.cells):
            target_cell = target_table.cell(row_idx, col_idx)
            target_cell.text = cell.text
            for paragraph in cell.paragraphs:
                new_paragraph = target_cell.paragraphs[0]
                new_paragraph.clear()
                for run in paragraph.runs:
                    new_run = new_paragraph.add_run(run.text)
                    new_run.bold = run.bold
                    new_run.italic = run.italic
                    new_run.underline = run.underline
                    new_run.font.name = run.font.name
                    new_run.font.size = run.font.size
                    new_run.font.color.rgb = run.font.color.rgb

def copy_footer(reference_section, target_section):
    """
    Copy footer from reference section to target section.
    """
    target_footer = target_section.footer

    # Clear existing footer content
    for element in target_footer.element.xpath('.//w:p'):
        element.getparent().remove(element)
    for element in target_footer.element.xpath('.//w:tbl'):
        element.getparent().remove(element)
    
    print("Copying footer...")

    for paragraph in reference_section.footer.paragraphs:
        new_paragraph = target_footer.add_paragraph()
        new_paragraph.text = paragraph.text
        new_paragraph.alignment = paragraph.alignment
        for run in paragraph.runs:
            new_run = new_paragraph.add_run(run.text)
            new_run.bold = run.bold
            new_run.italic = run.italic
            new_run.underline = run.underline
            new_run.font.name = run.font.name
            new_run.font.size = run.font.size
            new_run.font.color.rgb = run.font.color.rgb

    for table in reference_section.footer.tables:
        copy_table(table, target_footer)

def main():
    # Load the documents
    source_doc = Document('/home/exotic/Downloads/Reports/Combined_Report.docx')
    reference_doc = Document('/home/exotic/Downloads/PCT_HR_4002_F02 Training Evaluation Form_02.docx')

    # Assume the footer is in the first section of the reference document
    reference_section = reference_doc.sections[0]

    # Debug: Print footer contents of reference document
    print("Reference Footer:")
    for paragraph in reference_section.footer.paragraphs:
        print(paragraph.text)
    for table in reference_section.footer.tables:
        for row in table.rows:
            for cell in row.cells:
                print(cell.text)

    # Iterate through all sections in the source document
    for section in source_doc.sections:
        if not footer_exists(section):
            print(f"No footer found in section, copying footer.")
            copy_footer(reference_section, section)
        else:
            print(f"Footer already exists in section, skipping.")

    # Save the modified source document
    source_doc.save('/home/exotic/Downloads/Reports/modified_source_document.docx')
    print("Footer added and document saved as 'modified_source_document.docx'")

if __name__ == "__main__":
    main()




